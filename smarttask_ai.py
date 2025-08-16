#!/usr/bin/env python3
# SmartTask AI Assistant
# Author: Guillaume Lessard, iD01t Softwares
# Version: 1.0.0
#
# A premium AI-powered productivity suite designed to streamline your workflow.

# --- Core Python Libraries ---
import importlib.util
import subprocess
import sys
import os
import sqlite3
import json
import webbrowser
from datetime import datetime
import csv

# --- Dependency Management ---

def check_and_install_dependencies():
    """
    Checks for required packages and installs them if missing.
    This makes the script portable and easier to run by ensuring all
    necessary third-party libraries are available.
    """
    # A dictionary of package names for pip and their corresponding import names.
    required_packages = {
        "PyQt6": "PyQt6",
        "requests": "requests",
        "fpdf2": "fpdf",
        "python-docx": "docx",
        "PyPDF2": "PyPDF2"
    }

    print("--- SmartTask AI Assistant ---")
    print("Checking for required dependencies...")
    all_installed = True
    for package, import_name in required_packages.items():
        if importlib.util.find_spec(import_name) is None:
            all_installed = False
            print(f"'{package}' is not installed. Attempting to install...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", package],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
                )
                print(f"Successfully installed '{package}'.")
            except subprocess.CalledProcessError:
                print(f"ERROR: Failed to install '{package}'. Please install it manually.")
                sys.exit(1)

    if all_installed:
        print("All dependencies are satisfied.")
    print("-" * 30)

# --- Main Application ---

try:
    from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                                 QHBoxLayout, QListWidget, QTextEdit, QStatusBar,
                                 QPushButton, QLabel, QFrame, QStackedWidget,
                                 QLineEdit, QListWidgetItem, QFormLayout, QComboBox,
                                 QMessageBox, QFileDialog, QDateEdit)
    from PyQt6.QtGui import QIcon
    from PyQt6.QtCore import Qt, QSize, QDate
    from docx import Document
    from PyPDF2 import PdfReader
    from fpdf import FPDF
    import requests
except ImportError:
    check_and_install_dependencies()
    print("\nDependencies installed. Please restart the application.")
    sys.exit(0)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.db_conn = None
        self.init_database()
        self.setWindowTitle("SmartTask AI Assistant")
        self.setGeometry(100, 100, 1200, 800)
        if os.path.exists("icon.ico"): self.setWindowIcon(QIcon("icon.ico"))
        self.current_theme = 'dark'
        self.setup_ui()
        self.apply_theme(self.current_theme)
        self.load_api_keys_to_inputs()
        self.update_model_selector()
        self.update_status_bar()

    # --- Event Handlers ---
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls(): event.acceptProposedAction()
        else: event.ignore()

    def dropEvent(self, event):
        if not event.mimeData().hasUrls(): return
        if self.pages.currentWidget() != self.ai_chat_page:
            self.status_bar.showMessage("Please switch to AI Chat to drop files.", 4000); return
        filepath = event.mimeData().urls()[0].toLocalFile()
        self.handle_file_drop(filepath)

    def change_page(self): self.pages.setCurrentIndex(self.nav_list.currentRow())

    # --- UI Setup ---
    def setup_ui(self):
        self.central_widget = QWidget(); self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0); self.main_layout.setSpacing(0)
        self.setup_sidebar(); self.setup_main_content()
        self.status_bar = QStatusBar(); self.setStatusBar(self.status_bar)

    def setup_sidebar(self):
        self.sidebar = QWidget(); self.sidebar.setFixedWidth(200); self.sidebar_layout = QVBoxLayout(self.sidebar)
        self.sidebar_layout.setContentsMargins(10, 10, 10, 10); self.sidebar_layout.setSpacing(10); self.sidebar_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        title = QLabel("SmartTask AI"); title.setObjectName("SidebarTitle")
        self.nav_list = QListWidget(); self.nav_list.addItems(["Tasks", "AI Chat", "Settings"])
        self.nav_list.setCurrentRow(0); self.nav_list.itemSelectionChanged.connect(self.change_page)
        self.theme_button = QPushButton("Toggle Theme"); self.theme_button.clicked.connect(self.toggle_theme)
        self.sidebar_layout.addWidget(title); self.sidebar_layout.addWidget(self.nav_list); self.sidebar_layout.addStretch(); self.sidebar_layout.addWidget(self.theme_button)
        self.main_layout.addWidget(self.sidebar)

    def setup_main_content(self):
        self.pages = QStackedWidget(); self.main_layout.addWidget(self.pages)
        self.task_manager_page = self.create_task_manager_page(); self.ai_chat_page = self.create_ai_chat_page(); self.settings_page = self.create_settings_page()
        self.pages.addWidget(self.task_manager_page); self.pages.addWidget(self.ai_chat_page); self.pages.addWidget(self.settings_page)

    def create_task_manager_page(self):
        page = QWidget(); layout = QVBoxLayout(page); layout.setContentsMargins(20, 20, 20, 20)
        title = QLabel("Task Manager"); title.setObjectName("PageTitle"); self.task_list_widget = QListWidget(); self.load_tasks()
        input_layout = QHBoxLayout(); self.task_input = QLineEdit(); self.task_input.setPlaceholderText("Enter new task...")
        self.due_date_input = QDateEdit(self); self.due_date_input.setCalendarPopup(True); self.due_date_input.setDate(QDate.currentDate())
        add_button = QPushButton("Add Task"); add_button.clicked.connect(self.add_task)
        input_layout.addWidget(self.task_input, 1); input_layout.addWidget(self.due_date_input); input_layout.addWidget(add_button)
        delete_button = QPushButton("Delete Selected"); delete_button.clicked.connect(self.delete_task)
        layout.addWidget(title); layout.addWidget(self.task_list_widget); layout.addLayout(input_layout); layout.addWidget(delete_button, 0, Qt.AlignmentFlag.AlignRight)
        return page

    def create_ai_chat_page(self):
        page = QWidget(); layout = QVBoxLayout(page); layout.setContentsMargins(20, 20, 20, 20)
        title = QLabel("AI Chat"); title.setObjectName("PageTitle")
        model_area = QHBoxLayout(); model_label = QLabel("Select Model:"); self.model_selector = QComboBox()
        model_area.addWidget(model_label); model_area.addWidget(self.model_selector); model_area.addStretch()
        self.ai_output_display = QTextEdit(); self.ai_output_display.setReadOnly(True)
        self.command_input = QTextEdit(); self.command_input.setFixedHeight(120)
        button_layout = QHBoxLayout(); export_button = QPushButton("Export Chat"); export_button.clicked.connect(self.export_chat_history)
        send_button = QPushButton("Send to AI"); send_button.clicked.connect(self.handle_ai_request)
        button_layout.addStretch(); button_layout.addWidget(export_button); button_layout.addWidget(send_button)
        layout.addWidget(title); layout.addLayout(model_area); layout.addWidget(self.ai_output_display, 1); layout.addWidget(self.command_input); layout.addLayout(button_layout)
        return page

    def create_settings_page(self):
        page = QWidget(); layout = QVBoxLayout(page); layout.setContentsMargins(20, 20, 20, 20); layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        api_title = QLabel("API Key Management"); api_title.setObjectName("PageTitle"); api_form = QFormLayout(); self.api_key_inputs = {}
        for service in ["OpenAI", "Claude", "Gemini", "Custom Endpoint"]:
            self.api_key_inputs[service] = QLineEdit(); self.api_key_inputs[service].setEchoMode(QLineEdit.EchoMode.Password)
            api_form.addRow(f"<b>{service} API Key:</b>", self.api_key_inputs[service])
        save_api_button = QPushButton("Save API Keys"); save_api_button.clicked.connect(self.save_api_keys)
        license_title = QLabel("License & Monetization"); license_title.setObjectName("PageTitle"); license_form = QFormLayout()
        self.license_input = QLineEdit(); self.license_input.setPlaceholderText("Enter license key for Pro mode")
        license_form.addRow("<b>License Key:</b>", self.license_input)
        activate_button = QPushButton("Activate Pro"); activate_button.clicked.connect(self.activate_pro_license)
        promo_layout = QHBoxLayout(); gumroad_button = QPushButton("Get Pro on Gumroad"); gumroad_button.clicked.connect(lambda: webbrowser.open("https://gumroad.com/"))
        kofi_button = QPushButton("Support on Ko-fi"); kofi_button.clicked.connect(lambda: webbrowser.open("https://ko-fi.com/"))
        promo_layout.addWidget(gumroad_button); promo_layout.addWidget(kofi_button)
        layout.addWidget(api_title); layout.addLayout(api_form); layout.addWidget(save_api_button); layout.addSpacing(40)
        layout.addWidget(license_title); layout.addLayout(license_form); layout.addWidget(activate_button); layout.addSpacing(20); layout.addLayout(promo_layout)
        return page

    # --- Core Logic ---
    def init_database(self):
        self.db_conn = sqlite3.connect("smarttask.db");
        self.update_database_schema()
        cursor = self.db_conn.cursor()
        cursor.execute("CREATE TABLE IF NOT EXISTS api_keys (service TEXT PRIMARY KEY, api_key TEXT)")
        cursor.execute("CREATE TABLE IF NOT EXISTS app_config (key TEXT PRIMARY KEY, value TEXT)")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('license_status', 'UNLICENSED')")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('query_count', '0')")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('last_query_reset', ?)", (datetime.now().strftime('%Y-%m'),))
        self.db_conn.commit()

    def update_database_schema(self):
        cursor = self.db_conn.cursor()
        cursor.execute("CREATE TABLE IF NOT EXISTS tasks (id INTEGER PRIMARY KEY, description TEXT, status TEXT DEFAULT 'pending')")
        cursor.execute("PRAGMA table_info(tasks)")
        columns = [info[1] for info in cursor.fetchall()]
        if 'due_date' not in columns:
            cursor.execute("ALTER TABLE tasks ADD COLUMN due_date TEXT")
        self.db_conn.commit()

    def get_config(self, key):
        cursor = self.db_conn.cursor(); cursor.execute("SELECT value FROM app_config WHERE key=?", (key,)); return (res[0] if (res := cursor.fetchone()) else None)
    def set_config(self, key, value):
        cursor = self.db_conn.cursor(); cursor.execute("INSERT OR REPLACE INTO app_config VALUES (?, ?)", (key, value)); self.db_conn.commit()
    def update_status_bar(self):
        status = self.get_config('license_status')
        if status == 'PRO': self.status_bar.showMessage("Pro Version | Unlimited Queries")
        else: self.status_bar.showMessage(f"Free Version | Queries this month: {self.get_config('query_count')}/20")
    def activate_pro_license(self):
        key = self.license_input.text().strip().upper()
        if key.startswith("SMARTTASK-") and len(key) > 15:
            self.set_config('license_status', 'PRO'); QMessageBox.information(self, "Success", "Pro license activated!"); self.update_status_bar()
        else: QMessageBox.warning(self, "Error", "Invalid license key format.")
    def handle_ai_request(self):
        if self.get_config('license_status') != 'PRO':
            current_month = datetime.now().strftime('%Y-%m'); last_reset = self.get_config('last_query_reset')
            if current_month != last_reset: self.set_config('query_count', '0'); self.set_config('last_query_reset', current_month)
            query_count = int(self.get_config('query_count'))
            if query_count >= 20: QMessageBox.warning(self, "Limit Reached", "Query limit (20) reached. Please upgrade to Pro."); return
            self.set_config('query_count', str(query_count + 1)); self.update_status_bar()
        model = self.model_selector.currentText(); prompt = self.command_input.toPlainText().strip()
        if "No API" in model: self.status_bar.showMessage("Select a model in Settings.", 3000); return
        if not prompt: self.status_bar.showMessage("Please enter a prompt.", 3000); return
        cursor = self.db_conn.cursor(); cursor.execute("SELECT api_key FROM api_keys WHERE service=?", (model,)); res = cursor.fetchone()
        if not res: self.status_bar.showMessage(f"No API key for {model}.", 3000); return
        self.ai_output_display.append(f"<b>You:</b> {prompt}"); QApplication.processEvents()
        response = self.query_llm(model, prompt, res[0])
        self.ai_output_display.append(f"<b>{model}:</b> {response}"); self.command_input.clear()
    def load_tasks(self):
        self.task_list_widget.clear(); cursor = self.db_conn.cursor()
        for row in cursor.execute("SELECT id, description, due_date FROM tasks WHERE status='pending' ORDER BY due_date ASC, id DESC"):
            task_id, description, due_date = row
            display_text = f"{description}"
            if due_date: display_text += f"  (Due: {due_date})"
            item = QListWidgetItem(display_text); item.setData(Qt.ItemDataRole.UserRole, task_id); self.task_list_widget.addItem(item)
    def add_task(self):
        desc = self.task_input.text().strip()
        due_date = self.due_date_input.date().toString("yyyy-MM-dd")
        if desc:
            cursor = self.db_conn.cursor(); cursor.execute("INSERT INTO tasks (description, due_date) VALUES (?, ?)", (desc, due_date))
            self.db_conn.commit(); self.task_input.clear(); self.load_tasks(); self.status_bar.showMessage(f"Task added.", 2000)
    def delete_task(self):
        item = self.task_list_widget.currentItem()
        if item: task_id = item.data(Qt.ItemDataRole.UserRole); cursor = self.db_conn.cursor(); cursor.execute("DELETE FROM tasks WHERE id=?", (task_id,)); self.db_conn.commit(); self.load_tasks(); self.status_bar.showMessage("Task deleted.", 2000)
    def handle_file_drop(self, filepath):
        try:
            self.status_bar.showMessage(f"Analyzing {os.path.basename(filepath)}...", 3000)
            content = self.extract_text_from_file(filepath)
            if content:
                header = f"--- Context from {os.path.basename(filepath)} ---\n"; footer = "\n--- End of Context ---\n"
                self.command_input.setPlainText(header + content + footer); self.status_bar.showMessage("File content loaded.", 4000)
            else: QMessageBox.warning(self, "File Error", "Could not extract text from the file or file is empty.")
        except Exception as e: QMessageBox.critical(self, "File Read Error", f"Failed to process file: {e}")
    def extract_text_from_file(self, filepath):
        _, extension = os.path.splitext(filepath); extension = extension.lower()
        if extension == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f: return f.read()
        elif extension == '.csv':
            with open(filepath, 'r', newline='', encoding='utf-8', errors='ignore') as f: return "\n".join([",".join(row) for row in csv.reader(f)])
        elif extension == '.docx': return "\n".join([para.text for para in Document(filepath).paragraphs])
        elif extension == '.pdf': return "\n".join([page.extract_text() or "" for page in PdfReader(filepath).pages])
        else: self.status_bar.showMessage(f"Unsupported file type: {extension}", 3000); return None
    def load_api_keys_to_inputs(self):
        cursor = self.db_conn.cursor()
        for row in cursor.execute("SELECT service, api_key FROM api_keys"):
            if row[0] in self.api_key_inputs: self.api_key_inputs[row[0]].setText(row[1])
    def save_api_keys(self):
        cursor = self.db_conn.cursor()
        for service, field in self.api_key_inputs.items():
            key = field.text().strip()
            if key: cursor.execute("INSERT OR REPLACE INTO api_keys VALUES (?, ?)", (service, key))
            else: cursor.execute("DELETE FROM api_keys WHERE service=?", (service,))
        self.db_conn.commit(); self.status_bar.showMessage("API keys saved.", 3000); self.update_model_selector()
    def update_model_selector(self):
        self.model_selector.clear(); cursor = self.db_conn.cursor()
        keys = [row[0] for row in cursor.execute("SELECT service FROM api_keys")]
        if keys: self.model_selector.addItems(keys); self.model_selector.setEnabled(True)
        else: self.model_selector.addItem("No API Keys Set"); self.model_selector.setEnabled(False)
    def query_llm(self, model_name, prompt, api_key):
        """Dispatches the AI query to the appropriate function based on the model name."""
        try:
            if model_name == "OpenAI": return self.query_openai(prompt, api_key)
            elif model_name == "Claude": return self.query_claude(prompt, api_key)
            elif model_name == "Gemini": return self.query_gemini(prompt, api_key)
            elif model_name == "Custom Endpoint": return self.query_custom(prompt, api_key)
            else: return f"Model '{model_name}' not implemented yet."
        except requests.exceptions.RequestException as e: return f"<font color='red'><b>Network Error:</b> {e}</font>"
        except Exception as e: return f"<font color='red'><b>Error:</b> {e}</font>"

    def query_openai(self, prompt, api_key):
        """Sends a request to the OpenAI Chat Completions API."""
        endpoint = "https://api.openai.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        data = {"model": "gpt-3.5-turbo", "messages": [{"role": "user", "content": prompt}]}
        response = requests.post(endpoint, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']

    def query_claude(self, prompt, api_key):
        """Sends a request to the Anthropic Claude API."""
        endpoint = "https://api.anthropic.com/v1/messages"
        headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "Content-Type": "application/json"}
        data = {
            "model": "claude-3-sonnet-20240229", "max_tokens": 4096,
            "messages": [{"role": "user", "content": prompt}]
        }
        response = requests.post(endpoint, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        return response.json()['content'][0]['text']

    def query_gemini(self, prompt, api_key):
        """Sends a request to the Google Gemini API."""
        model = "gemini-1.5-flash-latest"
        endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
        headers = {"Content-Type": "application/json"}
        params = {"key": api_key}
        data = {"contents": [{"parts": [{"text": prompt}]}]}
        response = requests.post(endpoint, params=params, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        return response.json()['candidates'][0]['content']['parts'][0]['text']

    def query_custom(self, prompt, url):
        """Sends a request to a user-defined custom endpoint."""
        try:
            data = {"prompt": prompt}
            response = requests.post(url, json=data, timeout=30)
            response.raise_for_status()
            json_response = response.json()
            return json_response.get("response", json_response.get("text", str(json_response)))
        except requests.exceptions.RequestException as e:
            return f"<font color='red'>Error connecting to custom endpoint: {e}</font>"
    def toggle_theme(self):
        self.current_theme = 'light' if self.current_theme == 'dark' else 'dark'; self.apply_theme(self.current_theme)
    def apply_theme(self, theme):
        base = "QPushButton { border-radius: 5px; padding: 8px; font-weight: bold; } QListWidget, QTextEdit, QLineEdit, QDateEdit { border-radius: 5px; padding: 5px; } #SidebarTitle { font-size: 18pt; font-weight: bold; } #PageTitle { font-size: 14pt; font-weight: bold; margin-bottom: 10px; }"
        light = base + "QWidget { background-color: #f0f2f5; color: #333; } QTextEdit, QLineEdit, QListWidget, QDateEdit { background-color: #fff; border: 1px solid #d9d9d9; } QPushButton { background-color: #007bff; color: white; border: none; } QPushButton:hover { background-color: #0056b3; } QStatusBar { background-color: #e9ecef; } #sidebar { background-color: #fff; }"
        dark = base + "QWidget { background-color: #1c1c1e; color: #f0f0f0; } QTextEdit, QLineEdit, QListWidget, QDateEdit { background-color: #2c2c2e; border: 1px solid #444; } QPushButton { background-color: #0a84ff; color: white; border: none; } QPushButton:hover { background-color: #0060df; } QStatusBar { background-color: #2c2c2e; } #sidebar { background-color: #232325; }"
        self.setStyleSheet(dark if theme == 'dark' else light)
        self.sidebar.setObjectName("sidebar"); self.style().unpolish(self); self.style().polish(self)
    def export_chat_history(self):
        content = self.ai_output_display.toPlainText()
        if not content.strip(): QMessageBox.information(self, "Export Empty", "There is no chat history to export."); return
        file_path, selected_filter = QFileDialog.getSaveFileName(self, "Export Chat History", "", "PDF (*.pdf);;Word Document (*.docx);;Markdown (*.md)")
        if not file_path: return
        try:
            if selected_filter == "PDF (*.pdf)": self.export_to_pdf(content, file_path)
            elif selected_filter == "Word Document (*.docx)": self.export_to_docx(content, file_path)
            elif selected_filter == "Markdown (*.md)": self.export_to_md(content, file_path)
            self.status_bar.showMessage(f"Chat exported to {os.path.basename(file_path)}", 4000)
        except Exception as e: QMessageBox.critical(self, "Export Error", f"Failed to export file.\nError: {e}")
    def export_to_md(self, content, filepath):
        with open(filepath, 'w', encoding='utf-8') as f: f.write(content)
    def export_to_docx(self, content, filepath):
        doc = Document(); doc.add_paragraph(content); doc.save(filepath)
    def export_to_pdf(self, content, filepath):
        pdf = FPDF(font_cache_dir=True); pdf.add_page()
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf'); pdf.set_font('DejaVu', '', 12)
        pdf.multi_cell(0, 10, content); pdf.output(filepath)

if __name__ == '__main__':
    check_and_install_dependencies()
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
