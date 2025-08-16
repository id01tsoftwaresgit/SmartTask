#!/usr/bin/env python3
# SmartTask AI Assistant
# Version: 1.0.0
# Author: Guillaume Lessard, iD01t Softwares
# Website: https://www.id01t.ca
# License: MIT

# --- Core Python Libraries ---
import importlib.util
import subprocess
import sys
import os
import sqlite3
import webbrowser
import csv
from datetime import datetime

def check_and_install_dependencies():
    """
    Checks for required packages and installs them if missing.
    This makes the script portable and easier to run.
    """
    required_packages = {
        "PyQt6": "PyQt6", "requests": "requests", "fpdf2": "fpdf",
        "python-docx": "docx", "PyPDF2": "PyPDF2"
    }
    print("--- SmartTask AI Assistant: Dependency Check ---")
    all_installed = True
    for package, import_name in required_packages.items():
        if importlib.util.find_spec(import_name) is None:
            all_installed = False
            print(f"Installing missing package: {package}...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", package],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
                )
                print(f" -> Successfully installed '{package}'.")
            except subprocess.CalledProcessError:
                print(f" -> ERROR: Failed to install '{package}'. Please install it manually.")
                sys.exit(1)
    if all_installed:
        print("All dependencies are satisfied.")
    print("-" * 45)

# --- Main Application ---
try:
    from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                                 QTabWidget, QPushButton, QToolBar, QLabel,
                                 QListWidget, QLineEdit, QHBoxLayout, QListWidgetItem,
                                 QDateTimeEdit, QFormLayout, QTextEdit, QComboBox,
                                 QMessageBox, QStatusBar, QFileDialog)
    from PyQt6.QtGui import QIcon, QColor, QBrush
    from PyQt6.QtCore import QSize, QDateTime, Qt
    import requests
    from docx import Document
    from PyPDF2 import PdfReader
    from fpdf import FPDF
except ImportError:
    check_and_install_dependencies()
    print("\nDependencies installed. Please restart the application.")
    sys.exit(0)

class MainWindow(QMainWindow):
    """
    Main application window for the SmartTask AI Assistant.
    This class orchestrates the entire application, from UI setup to backend logic.
    """
    def __init__(self):
        """Initializes the application window, database, and UI."""
        super().__init__()
        self.setAcceptDrops(True)
        self.db_conn = None
        self.init_database()

        self.setWindowTitle("SmartTask AI Assistant")
        self.setWindowIcon(QIcon("icon.ico"))
        self.setGeometry(100, 100, 1200, 800)

        self.current_theme = 'dark'
        self.setup_ui()
        self.apply_theme()

        self.load_api_keys()
        self.update_model_selector()
        self.update_status_bar()

    # --- Event Handlers ---
    def dragEnterEvent(self, event):
        """Accepts drag events if they contain file URLs."""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        """Handles dropped files, extracting text for AI context."""
        if not event.mimeData().hasUrls() or self.tabs.currentWidget() != self.ai_chat_tab:
            return
        filepath = event.mimeData().urls()[0].toLocalFile()
        try:
            content = self.extract_text_from_file(filepath)
            if content:
                self.command_input.setPlainText(f"--- Context from {os.path.basename(filepath)} ---\n{content}\n--- End of Context ---\n")
                self.status_bar.showMessage(f"Loaded context from {os.path.basename(filepath)}", 4000)
            else:
                QMessageBox.warning(self, "File Error", "Could not extract text from the file.")
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Failed to process file: {e}")

    # --- UI Setup ---
    def setup_ui(self):
        """Initializes and arranges all UI components."""
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        # Create and add tabs
        self.task_tab = self.create_task_manager_tab()
        self.ai_chat_tab = self.create_ai_chat_tab()
        self.settings_tab = self.create_settings_tab()
        self.tabs.addTab(self.task_tab, "Tasks")
        self.tabs.addTab(QWidget(), "Documents") # Placeholder
        self.tabs.addTab(self.ai_chat_tab, "AI Chat")
        self.tabs.addTab(QWidget(), "Email") # Placeholder
        self.tabs.addTab(self.settings_tab, "Settings")
        # Toolbar
        toolbar = QToolBar("Main Toolbar")
        self.addToolBar(toolbar)
        self.theme_button = QPushButton("Toggle Theme")
        self.theme_button.clicked.connect(self.toggle_theme)
        toolbar.addWidget(self.theme_button)
        # Status Bar
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

    def create_task_manager_tab(self):
        """Creates the UI for the Task Manager tab."""
        page = QWidget()
        layout = QVBoxLayout(page)
        self.task_list_widget = QListWidget()
        self.load_tasks()
        # Input controls
        input_layout = QHBoxLayout()
        self.task_input = QLineEdit(placeholderText="Enter a new task...")
        self.due_date_input = QDateTimeEdit(QDateTime.currentDateTime())
        self.due_date_input.setCalendarPopup(True)
        add_button = QPushButton("Add Task")
        add_button.clicked.connect(self.add_task)
        input_layout.addWidget(self.task_input, 1)
        input_layout.addWidget(self.due_date_input)
        input_layout.addWidget(add_button)
        delete_button = QPushButton("Delete Selected Task")
        delete_button.clicked.connect(self.delete_task)
        # Assemble layout
        layout.addWidget(self.task_list_widget)
        layout.addLayout(input_layout)
        layout.addWidget(delete_button, alignment=Qt.AlignmentFlag.AlignRight)
        return page

    def create_ai_chat_tab(self):
        """Creates the UI for the AI Chat tab."""
        page = QWidget()
        layout = QVBoxLayout(page)
        # Model selection dropdown
        model_area = QHBoxLayout()
        self.model_selector = QComboBox()
        model_area.addWidget(QLabel("Select Model:"))
        model_area.addWidget(self.model_selector)
        model_area.addStretch()
        # Main chat panels
        self.ai_output_display = QTextEdit(readOnly=True)
        self.command_input = QTextEdit(placeholderText="Enter your command or question...")
        self.command_input.setFixedHeight(100)
        # Action buttons
        button_layout = QHBoxLayout()
        export_button = QPushButton("Export Chat")
        export_button.clicked.connect(self.export_chat_history)
        send_button = QPushButton("Send to AI")
        send_button.clicked.connect(self.handle_ai_request)
        button_layout.addStretch()
        button_layout.addWidget(export_button)
        button_layout.addWidget(send_button)
        # Assemble layout
        layout.addLayout(model_area)
        layout.addWidget(self.ai_output_display, 1)
        layout.addWidget(self.command_input)
        layout.addLayout(button_layout)
        return page

    def create_settings_tab(self):
        """Creates the UI for the Settings tab."""
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        # API Keys Section
        api_form = QFormLayout()
        self.api_key_inputs = {}
        for service in ["OpenAI", "Claude", "Gemini", "Custom Endpoint"]:
            self.api_key_inputs[service] = QLineEdit(echoMode=QLineEdit.EchoMode.Password)
            api_form.addRow(f"<b>{service} API Key:</b>", self.api_key_inputs[service])
        save_api_button = QPushButton("Save API Keys")
        save_api_button.clicked.connect(self.save_api_keys)
        # License & Monetization Section
        license_form = QFormLayout()
        self.license_input = QLineEdit(placeholderText="Enter license key to unlock Pro")
        license_form.addRow("<b>License Key:</b>", self.license_input)
        activate_button = QPushButton("Activate Pro")
        activate_button.clicked.connect(self.activate_pro_license)
        # Promo Buttons
        promo_layout = QHBoxLayout()
        gumroad_button = QPushButton("Get Pro on Gumroad")
        gumroad_button.clicked.connect(lambda: webbrowser.open("https://gumroad.com/"))
        kofi_button = QPushButton("Support on Ko-fi")
        kofi_button.clicked.connect(lambda: webbrowser.open("https://ko-fi.com/"))
        promo_layout.addWidget(gumroad_button)
        promo_layout.addWidget(kofi_button)
        # Assemble layout
        layout.addLayout(api_form)
        layout.addWidget(save_api_button, alignment=Qt.AlignmentFlag.AlignLeft)
        layout.addSpacing(20)
        layout.addLayout(license_form)
        layout.addWidget(activate_button, alignment=Qt.AlignmentFlag.AlignLeft)
        layout.addSpacing(20)
        layout.addStretch()
        layout.addLayout(promo_layout)
        return page

    # --- Database and Config Logic ---
    def init_database(self):
        """Initializes the database and creates tables if they don't exist."""
        self.db_conn = sqlite3.connect("smarttask.db")
        cursor = self.db_conn.cursor()
        cursor.execute("CREATE TABLE IF NOT EXISTS tasks (id INTEGER PRIMARY KEY, description TEXT, due_date TEXT, status TEXT DEFAULT 'pending')")
        cursor.execute("CREATE TABLE IF NOT EXISTS api_keys (service TEXT PRIMARY KEY, api_key TEXT NOT NULL)")
        cursor.execute("CREATE TABLE IF NOT EXISTS app_config (key TEXT PRIMARY KEY, value TEXT)")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('license_status', 'UNLICENSED')")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('query_count', '0')")
        cursor.execute("INSERT OR IGNORE INTO app_config VALUES ('last_query_reset', ?)", (datetime.now().strftime('%Y-%m'),))
        self.db_conn.commit()
    def get_config(self, key):
        """Gets a value from the app_config table."""
        cursor = self.db_conn.cursor()
        cursor.execute("SELECT value FROM app_config WHERE key=?", (key,))
        return (res[0] if (res := cursor.fetchone()) else None)
    def set_config(self, key, value):
        """Sets a value in the app_config table."""
        cursor = self.db_conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO app_config VALUES (?, ?)", (key, value))
        self.db_conn.commit()

    # --- Feature Logic ---
    def update_status_bar(self):
        """Updates the status bar text with the current license status."""
        status = self.get_config('license_status')
        if status == 'PRO':
            self.status_bar.showMessage("Pro Version | Unlimited Queries")
        else:
            self.status_bar.showMessage(f"Free Version | Queries this month: {self.get_config('query_count')}/20")
    def activate_pro_license(self):
        """Validates and activates a Pro license."""
        key = self.license_input.text().strip().upper()
        if key.startswith("SMARTTASK-") and len(key) > 15:
            self.set_config('license_status', 'PRO')
            QMessageBox.information(self, "Success", "Pro license activated!")
            self.update_status_bar()
        else:
            QMessageBox.warning(self, "Error", "Invalid license key format.")
    def handle_ai_request(self):
        """Handles the entire AI request lifecycle, including license checks."""
        if self.get_config('license_status') != 'PRO':
            current_month = datetime.now().strftime('%Y-%m')
            last_reset = self.get_config('last_query_reset')
            if current_month != last_reset:
                self.set_config('query_count', '0')
                self.set_config('last_query_reset', current_month)
            query_count = int(self.get_config('query_count'))
            if query_count >= 20:
                QMessageBox.warning(self, "Limit Reached", "You have reached your monthly query limit (20). Please upgrade to Pro.")
                return
            self.set_config('query_count', str(query_count + 1))
            self.update_status_bar()
        model = self.model_selector.currentText()
        prompt = self.command_input.toPlainText().strip()
        if not prompt or "No API" in model: return
        cursor = self.db_conn.cursor()
        cursor.execute("SELECT api_key FROM api_keys WHERE service=?", (model,))
        result = cursor.fetchone()
        if not result:
            QMessageBox.warning(self, "API Key Missing", f"API key for {model} not found. Please add it in Settings.")
            return
        self.ai_output_display.append(f"<b>You:</b> {prompt}")
        QApplication.processEvents()
        response = self.query_llm(model, prompt, result[0])
        self.ai_output_display.append(f"<b>{model}:</b> {response}")
        self.command_input.clear()
    def query_llm(self, model, prompt, key):
        """Dispatches the query to the correct LLM API."""
        try:
            if model == "OpenAI": return self.query_openai(prompt, key)
            if model == "Claude": return f"[Placeholder: Would query Claude with prompt: '{prompt[:30]}...']"
            if model == "Gemini": return f"[Placeholder: Would query Gemini with prompt: '{prompt[:30]}...']"
            return "Model not implemented."
        except Exception as e:
            return f"<font color='red'>Error: {e}</font>"
    def query_openai(self, prompt, key):
        """Sends a request to the OpenAI API."""
        response = requests.post("https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {key}"},
            json={"model": "gpt-3.5-turbo", "messages": [{"role": "user", "content": prompt}]})
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    def save_api_keys(self):
        """Saves API keys to the database."""
        cursor = self.db_conn.cursor()
        for service, field in self.api_key_inputs.items():
            key = field.text().strip()
            if key: cursor.execute("INSERT OR REPLACE INTO api_keys VALUES (?, ?)", (service, key))
            else: cursor.execute("DELETE FROM api_keys WHERE service = ?", (service,))
        self.db_conn.commit()
        self.update_model_selector()
        self.status_bar.showMessage("API Keys saved successfully.", 3000)
    def load_api_keys(self):
        """Loads API keys from the database."""
        cursor = self.db_conn.cursor()
        for row in cursor.execute("SELECT service, api_key FROM api_keys"):
            if row[0] in self.api_key_inputs: self.api_key_inputs[row[0]].setText(row[1])
    def update_model_selector(self):
        """Updates the model selector dropdown with available keys."""
        self.model_selector.clear()
        keys = [row[0] for row in self.db_conn.cursor().execute("SELECT service FROM api_keys")]
        if keys: self.model_selector.addItems(keys)
        else: self.model_selector.addItem("No API Keys Set")
    def load_tasks(self):
        """Loads tasks and applies reminder colors."""
        self.task_list_widget.clear()
        now = datetime.now()
        for row in self.db_conn.cursor().execute("SELECT id, description, due_date FROM tasks WHERE status='pending' ORDER BY due_date ASC"):
            task_id, description, due_date_str = row
            display_text = f"{description}"
            if due_date_str:
                due_date = datetime.fromisoformat(due_date_str)
                display_text += f" (Due: {due_date.strftime('%Y-%m-%d %H:%M')})"
            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, task_id)
            if due_date_str:
                if due_date < now: item.setForeground(QBrush(QColor("#e57373"))) # Red
                elif (due_date - now).days < 1: item.setForeground(QBrush(QColor("#ffb74d"))) # Orange
            self.task_list_widget.addItem(item)
    def add_task(self):
        """Adds a new task to the database."""
        desc = self.task_input.text().strip()
        if not desc: return
        due_date = self.due_date_input.dateTime().toString(Qt.DateFormat.ISODate)
        self.db_conn.cursor().execute("INSERT INTO tasks (description, due_date) VALUES (?, ?)", (desc, due_date))
        self.db_conn.commit()
        self.task_input.clear()
        self.load_tasks()
    def delete_task(self):
        """Deletes the selected task."""
        selected_item = self.task_list_widget.currentItem()
        if not selected_item: return
        task_id = selected_item.data(Qt.ItemDataRole.UserRole)
        self.db_conn.cursor().execute("DELETE FROM tasks WHERE id=?", (task_id,))
        self.db_conn.commit()
        self.load_tasks()
    def toggle_theme(self):
        """Switches the application theme."""
        self.current_theme = 'light' if self.current_theme == 'dark' else 'dark'
        self.apply_theme()
    def apply_theme(self):
        """Applies the current theme's stylesheet."""
        dark_stylesheet = "QWidget{background-color:#2b2b2b;color:#f0f0f0}QTabWidget::pane{border-top:2px solid #3c3f41}QTabBar::tab{background:#2b2b2b;border:1px solid #444;padding:10px}QTabBar::tab:selected{background:#3c3f41}QPushButton{background-color:#555;border:1px solid #666;padding:5px}QPushButton:hover{background-color:#666}QToolBar{background-color:#3c3f41;border:none}QLineEdit,QDateTimeEdit,QTextEdit{border:1px solid #444;padding:5px;background-color:#3c3f41}QListWidget{border:1px solid #444}"
        light_stylesheet = "QWidget{background-color:#f0f0f0;color:#333}QTabWidget::pane{border-top:2px solid #d9d9d9}QTabBar::tab{background:#f0f0f0;border:1px solid #ccc;padding:10px}QTabBar::tab:selected{background:#fff}QPushButton{background-color:#e0e0e0;border:1px solid #ccc;padding:5px}QPushButton:hover{background-color:#d4d4d4}QToolBar{background-color:#e9e9e9;border:none}QLineEdit,QDateTimeEdit,QTextEdit{border:1px solid #ccc;padding:5px;background-color:#fff}QListWidget{border:1px solid #ccc}"
        self.setStyleSheet(dark_stylesheet if self.current_theme == 'dark' else light_stylesheet)
    def extract_text_from_file(self, filepath):
        """Extracts text from a given file path."""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f: return f.read()
        elif ext == '.csv':
            with open(filepath, 'r', newline='', encoding='utf-8') as f: return "\n".join([",".join(row) for row in csv.reader(f)])
        elif ext == '.docx': return "\n".join([p.text for p in Document(filepath).paragraphs])
        elif ext == '.pdf': return "\n".join([page.extract_text() or "" for page in PdfReader(filepath).pages])
        return None
    def export_chat_history(self):
        """Exports the chat history to a user-selected file."""
        content = self.ai_output_display.toPlainText()
        if not content.strip(): QMessageBox.information(self, "Export Empty", "There is no chat history to export."); return
        file_path, selected_filter = QFileDialog.getSaveFileName(self, "Export Chat History", "", "PDF (*.pdf);;Word Document (*.docx);;Markdown (*.md)")
        if not file_path: return
        try:
            if selected_filter.startswith("PDF"): self.export_to_pdf(content, file_path)
            elif selected_filter.startswith("Word"): self.export_to_docx(content, file_path)
            elif selected_filter.startswith("Markdown"): self.export_to_md(content, file_path)
            self.status_bar.showMessage(f"Chat exported to {os.path.basename(file_path)}", 4000)
        except Exception as e: QMessageBox.critical(self, "Export Error", f"Failed to export file.\nError: {e}")
    def export_to_md(self, content, filepath):
        """Saves content to a Markdown file."""
        with open(filepath, 'w', encoding='utf-8') as f: f.write(content)
    def export_to_docx(self, content, filepath):
        """Saves content to a DOCX file."""
        doc = Document(); doc.add_paragraph(content); doc.save(filepath)
    def export_to_pdf(self, content, filepath):
        """Saves content to a PDF file with Unicode support."""
        pdf = FPDF(font_cache_dir=True); pdf.add_page()
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf'); pdf.set_font('DejaVu', '', 12)
        pdf.multi_cell(0, 10, content); pdf.output(filepath)

if __name__ == '__main__':
    check_and_install_dependencies()
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
